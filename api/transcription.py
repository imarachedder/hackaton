from pyannote.audio import Pipeline
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from pydub import AudioSegment
from dotenv import load_dotenv
import requests
import json
import re
import os
import whisper
import pathlib


# Load environment variables from .env file
load_dotenv()

"""
Глобальная проблема:
Pydub и другие фреймфорки по работе с аудио файлами не могут работать корректно с форматом .ogg 
"""
def convert_to_wav(input_path: pathlib.Path, output_path: pathlib.Path) -> None:
    try:
        audio_format = input_path.suffix[1:].lower()
        audio = AudioSegment.from_file(input_path, format=audio_format)
        audio.export(output_path, format="wav")
    except Exception as e:
        print(f"Error during conversion: {e}")

def transcribe_audio(file_path: pathlib.Path):
    """
    на своих мощностях выбрал пока base model
    также можно выбрать другие модели
    https://github.com/openai/whisper/blob/main/README.md
    :param file_path:
    :return:
    """

    model = whisper.load_model("medium")
    result = model.transcribe(file_path, language='ru')

    return result['text']

def diarize_audio(file_path: pathlib.Path):
    """
    Пытался вывести количество спикеров, но пока увы безуспешно
    :param file_path:
    :return:
    """
    pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization",
                                    use_auth_token="T")
    diarization = pipeline({'uri': 'filename', 'audio': file_path})
    return diarization

def generate_text_with_gigachat(prompt: str, api_key: str) -> str:
    """
    Для гигачата необходимы сертификаты минцифр для корректного запроса
    пока что работает на запросе по одному токену.
    Можно масштабировать по количеству токеном для многопоточности
    :param prompt:
    :param api_key:
    :return:
    """
    cert_path = 'russian_trusted_root_ca.cer'
    url = "https://gigachat.devices.sberbank.ru/api/v1/chat/completions"  # URL может отличаться, проверьте документацию GigaChat
    headers = {
        "Authorization": f"Bearer {api_key}",
        'Content-Type': 'application/json',
        'Accept': 'application/json'
    }
    payload = json.dumps({
        "model": "GigaChat",
        "messages": [
            {
                "role": "system",
                "content": "Ты профессиональный редактор и обработчик текстов. Отредактируй текст пользователя."
                           "Также выдели цель встречи. Расскажи тезисно, о чем там говорилось в пункте описание."
                           "Если в тексте больше одного человека, напиши количество спикеров"
            },
            {
                "role": "user",
                "content": prompt
            }
        ],
        "n": 1,
        "stream": False,
        "update_interval": 0
    })
    response = requests.post(url, headers=headers, data=payload, verify=cert_path)
    if response.status_code == 200:
        result = response.json()
        with open("text.txt", "w") as f:
            f.write(result.get('choices', [{}])[0].get('message', {}).get('content', ''))
        return result.get('choices', [{}])[0].get('message', {}).get('content', '')
    else:
        print(f"Ошибка: {response.status_code}")
        return ""

def create_document(transcription, summary):
    """

    :param transcription:
    :param summary:
    :param filename:
    :return:
    """
    files = ["raw.docx", "unformal.docx", "official.docx"]
    for file in files:
        doc = Document()
        doc.add_heading('Протокол совещания', 0)
        doc.add_heading('Количество спикеров:', level=1)
        doc.add_paragraph()
        doc.add_heading('Транскрипция:', level=1)
        doc.add_paragraph(transcription)
        doc.add_heading('Обработанный текст:', level=1)
        doc.add_paragraph(summary)
        tasks = extract_tasks(transcription)
        doc.add_heading('Поставленные задачи:', level=1)
        doc.add_paragraph(tasks)
        doc.save(file)


def create_pdf_documents(transcription, summary):
    files = ["raw.pdf", "unformal.pdf", "official.pdf"]

    for file in files:
        # Создание PDF-документа
        c = canvas.Canvas(file, pagesize=letter)
        width, height = letter
        y_position = height - 50  # Начальная позиция текста по оси Y

        # Заголовок документа
        c.setFont("Helvetica-Bold", 16)
        c.drawString(100, y_position, "Протокол совещания")
        y_position -= 30

        # Количество спикеров (если нужно добавить конкретные данные, можно передать параметр)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, y_position, "Количество спикеров:")
        y_position -= 20

        # Транскрипция
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, y_position, "Транскрипция:")
        y_position -= 20
        c.setFont("Helvetica", 10)
        text = transcription.split('\n')
        for line in text:
            c.drawString(100, y_position, line)
            y_position -= 15
            if y_position < 50:
                c.showPage()  # Переход на новую страницу
                y_position = height - 50

        # Обработанный текст
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, y_position, "Обработанный текст:")
        y_position -= 20
        c.setFont("Helvetica", 10)
        text = summary.split('\n')
        for line in text:
            c.drawString(100, y_position, line)
            y_position -= 15
            if y_position < 50:
                c.showPage()
                y_position = height - 50

        # Поставленные задачи
        tasks = extract_tasks(transcription)
        c.setFont("Helvetica-Bold", 12)
        c.drawString(100, y_position, "Поставленные задачи:")
        y_position -= 20
        c.setFont("Helvetica", 10)
        c.drawString(100, y_position, tasks)

        # Сохранение PDF-файла
        c.save()

def extract_tasks(text):
    tasks = re.findall(r'(?:поручено|задача)\s+(.*?)(?:срок выполнения\s+(\d{4}-\d{2}-\d{2}))?', text, re.DOTALL)
    return tasks

def get_token():
    url = "https://ngw.devices.sberbank.ru:9443/api/v2/oauth"
    RqUID = os.getenv('RQUID')
    authorization = os.getenv('AUTHORIZATION')
    # cert_path = os.getenv('CERT_PATH')
    cert_path = 'russian_trusted_root_ca.cer'

    payload = 'scope=GIGACHAT_API_PERS'
    headers = {
        'Content-Type': 'application/x-www-form-urlencoded',
        'Accept': 'application/json',
        'RqUID': RqUID,
        'Authorization': f'Basic {authorization}'
    }

    response = requests.request("POST", url, headers=headers, data=payload, verify=cert_path)
    return response.json().get("access_token")

async def transcription(audio_file: pathlib.Path):
    # Основной процесс
    api_key = get_token()
    convert_to_wav(audio_file, pathlib.Path("files")/"voice"/audio_file.with_suffix(".wav"))
    transcription = transcribe_audio(pathlib.Path("files")/"voice"/audio_file.with_suffix(".wav"))
    # diarization = diarize_audio(r"D:\projects\hack\train\Встреча 2.wav")

    summary = generate_text_with_gigachat(transcription, api_key)
    create_document(transcription, summary)
    create_pdf_documents(transcription, summary)
    return transcription
